"""
RQS (Requirement Sheet) Parser
Parses the RQS .docx file and extracts all fields into a dictionary.
"""

import os
import re
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def extract_images(doc, output_dir):
    """Extract all images from the document to output_dir.
    Returns a list of saved image paths in order of appearance."""
    os.makedirs(output_dir, exist_ok=True)
    image_paths = []
    img_idx = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            ext = os.path.splitext(rel.target_ref)[1]
            if not ext:
                ext = ".png"
            img_name = f"rqs_image_{img_idx}{ext}"
            img_path = os.path.join(output_dir, img_name)
            with open(img_path, "wb") as f:
                f.write(img_data)
            image_paths.append(img_path)
            img_idx += 1

    return image_paths


def _clean_text(text):
    """Clean whitespace from extracted text."""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text).strip()


def _extract_table_keyval(table):
    """Extract key-value pairs from a 2-column or 4-column table.
    Returns a dict of {key: value}."""
    data = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = _clean_text(cells[0].text)
            val = _clean_text(cells[1].text)
            if key and key != val:  # avoid merged header cells
                data[key] = val
        if len(cells) >= 4:
            key2 = _clean_text(cells[2].text)
            val2 = _clean_text(cells[3].text)
            if key2 and key2 != val2:
                data[key2] = val2
    return data


def _extract_table_rows(table):
    """Extract all rows as lists of cell texts."""
    rows = []
    for row in table.rows:
        row_data = [_clean_text(cell.text) for cell in row.cells]
        rows.append(row_data)
    return rows


def parse_rqs(docx_path, image_output_dir=None):
    """
    Parse an RQS document and return a dict with all extracted fields.

    Returns:
        dict with keys like:
            - project_no, task_no, project, test_coordinator, component, etc.
            - test_regulation, test_object, test_fixture, sled_pulse, etc.
            - type_of_dummy, direction_of_acceleration, test_setup, camera_setup
            - h_point_x, h_point_y, h_point_z (targets)
            - seat_back_angle, seat_cushion_angle
            - images: list of image paths
            - raw_tables: list of all table data for debugging
    """
    doc = DocxDocument(docx_path)
    result = {}
    raw_tables = []

    # Extract all tables
    for table in doc.tables:
        table_data = _extract_table_keyval(table)
        raw_tables.append(table_data)

        for key, val in table_data.items():
            # Header table fields
            if "Project No" in key:
                result["project_no"] = val
            elif key == "Task No." or key == "Task No.:":
                result["task_no"] = val
            elif key == "Project:" or key == "Project":
                result["project"] = val
            elif "Test coordinator" in key:
                result["test_coordinator"] = val
            elif "Component" in key:
                result["component"] = val
            elif "Representative" in key:
                result["representative"] = val
            elif "Customer" in key:
                result["customer"] = val

            # Sample info
            elif "Sample ID" in key:
                result["sample_id"] = val
            elif "Sample Content" in key:
                result["sample_content"] = val
            elif "Fixture/BIW" in key:
                result["fixture_biw"] = val
            elif "Part Level" in key:
                result["part_level"] = val
            elif "Additional Notes" in key:
                result["additional_notes"] = val
            elif "Sample availability" in key:
                result["sample_availability"] = val

            # Test criteria / requirements
            elif "Test Regulation" in key:
                result["test_regulation"] = val
            elif "Test Object" in key:
                result["test_object"] = val
            elif "Test Fixture" in key:
                result["test_fixture"] = val
            elif "Sled Pulse" in key:
                result["sled_pulse"] = val
                # Try to extract pulse IDs
                pulse_matches = re.findall(r'[JjSs]\d{5,}', val)
                if pulse_matches:
                    result["pulse_id"] = pulse_matches[0]
            elif "Direction of Acceleration" in key or "Direction" in key:
                result["direction_of_acceleration"] = val
            elif "Type of Dummy" in key or "Dummy type" in key:
                result["type_of_dummy"] = val
            elif "Dummy" in key and "%" in val:
                result["type_of_dummy"] = val

            # Test setup
            elif "Test Setup" in key:
                result["test_setup"] = val
            elif "Camera" in key:
                result["camera_setup"] = val

            # Seat positions
            elif "Seat Position" in key or "Seat position" in key:
                result["seat_position"] = val
            elif "Seat Back Angle" in key or "Back Angle" in key or "Backrest" in key:
                result["seat_back_angle"] = val
            elif "Seat Cushion Angle" in key or "Cushion Angle" in key:
                result["seat_cushion_angle"] = val
            elif "Head Restraint" in key or "Headrest" in key:
                result["head_restraint"] = val

    # Extract H-Point values from dummy setup tables
    for table in doc.tables:
        rows = _extract_table_rows(table)
        for i, row in enumerate(rows):
            row_text = " ".join(row).lower()
            if "h-point" in row_text or "h point" in row_text:
                for j, cell_text in enumerate(row):
                    if "X" in cell_text.upper() and j + 1 < len(row):
                        # Look for target value in same row or next column
                        pass
                # Try to get X, Y, Z from subsequent rows
                for k in range(i, min(i + 4, len(rows))):
                    r = rows[k]
                    for ci, c in enumerate(r):
                        c_upper = c.strip().upper()
                        if c_upper == "X" and ci + 1 < len(r):
                            result["h_point_x_target"] = r[ci + 1] if ci + 1 < len(r) else ""
                            if ci + 2 < len(r):
                                result["h_point_x_actual"] = r[ci + 2]
                        elif c_upper == "Y" and ci + 1 < len(r):
                            result["h_point_y_target"] = r[ci + 1] if ci + 1 < len(r) else ""
                            if ci + 2 < len(r):
                                result["h_point_y_actual"] = r[ci + 2]
                        elif c_upper == "Z" and ci + 1 < len(r):
                            result["h_point_z_target"] = r[ci + 1] if ci + 1 < len(r) else ""
                            if ci + 2 < len(r):
                                result["h_point_z_actual"] = r[ci + 2]

    # Extract test procedure, video analysis from paragraphs
    full_text_parts = []
    current_section = None
    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if not text:
            continue
        full_text_parts.append(text)

        lower = text.lower()
        if "test procedure" in lower:
            current_section = "test_procedure"
            result.setdefault("test_procedure", [])
        elif "video analysis" in lower or "video analyse" in lower:
            current_section = "video_analysis"
            result.setdefault("video_analysis", [])
        elif "evaluation" in lower and "method" in lower:
            current_section = "evaluation_method"
            result.setdefault("evaluation_method", [])
        elif current_section:
            result.setdefault(current_section, []).append(text)

    # Extract dates
    for table in doc.tables:
        table_data = _extract_table_keyval(table)
        for key, val in table_data.items():
            if "Start" in key and "Plan" in key:
                result["start_date"] = val
            elif "End" in key and "Plan" in key:
                result["end_date"] = val

    # Extract images
    if image_output_dir:
        result["images"] = extract_images(doc, image_output_dir)
    else:
        result["images"] = []

    result["raw_tables"] = raw_tables

    # Derive additional fields
    if "sled_pulse" in result and "pulse_id" not in result:
        m = re.search(r'[A-Z]\d{5,}', result["sled_pulse"])
        if m:
            result["pulse_id"] = m.group()

    # Extract Biltir number if present
    biltir_match = re.search(r'Biltir[:\s]*(\S+)', result.get("sled_pulse", ""))
    if biltir_match:
        result["biltir_number"] = biltir_match.group(1)

    return result


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python rqs_parser.py <rqs_file.docx>")
        sys.exit(1)

    data = parse_rqs(sys.argv[1], image_output_dir="./rqs_images")
    # Remove raw_tables for cleaner output
    display = {k: v for k, v in data.items() if k != "raw_tables"}
    print(json.dumps(display, indent=2, ensure_ascii=False, default=str))
