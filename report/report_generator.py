"""
Report Generator
Takes a template taslak .docx and fills it with data parsed from an RQS .docx.
Uses python-docx for find-and-replace to preserve all original formatting.
"""

import os
import re
import copy
from docx import Document as DocxDocument
from docx.shared import Inches
from report.rqs_parser import parse_rqs


def _replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting.
    Handles text split across multiple runs."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Simple case: text is in a single run
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # Complex case: text spans multiple runs - rebuild
    # Find start position in full text
    start_idx = full_text.index(old_text)
    end_idx = start_idx + len(old_text)

    # Map character positions to runs
    char_to_run = []
    for run_idx, run in enumerate(paragraph.runs):
        for _ in run.text:
            char_to_run.append(run_idx)

    if end_idx > len(char_to_run):
        return False

    start_run = char_to_run[start_idx]
    end_run = char_to_run[end_idx - 1]

    # Calculate offsets within runs
    start_offset = start_idx - sum(len(paragraph.runs[i].text) for i in range(start_run))
    end_offset = end_idx - sum(len(paragraph.runs[i].text) for i in range(end_run))

    # Replace in runs
    if start_run == end_run:
        run = paragraph.runs[start_run]
        run.text = run.text[:start_offset] + new_text + run.text[end_offset:]
    else:
        # Put new text in first run, clear middle runs, trim last run
        paragraph.runs[start_run].text = paragraph.runs[start_run].text[:start_offset] + new_text
        for i in range(start_run + 1, end_run):
            paragraph.runs[i].text = ""
        paragraph.runs[end_run].text = paragraph.runs[end_run].text[end_offset:]

    return True


def _replace_in_table(table, old_text, new_text):
    """Replace text in all cells of a table."""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if _replace_in_paragraph(paragraph, old_text, new_text):
                    replaced = True
            # Handle nested tables
            for nested_table in cell.tables:
                if _replace_in_table(nested_table, old_text, new_text):
                    replaced = True
    return replaced


def _replace_in_document(doc, old_text, new_text):
    """Replace text throughout the entire document."""
    replaced = False
    for paragraph in doc.paragraphs:
        if _replace_in_paragraph(paragraph, old_text, new_text):
            replaced = True
    for table in doc.tables:
        if _replace_in_table(table, old_text, new_text):
            replaced = True
    # Also check headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    if _replace_in_paragraph(paragraph, old_text, new_text):
                        replaced = True
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    if _replace_in_paragraph(paragraph, old_text, new_text):
                        replaced = True
    return replaced


def generate_report(template_path, rqs_data, output_path, replacements=None):
    """
    Generate a report by filling the template with RQS data.

    Args:
        template_path: Path to the taslak template .docx
        rqs_data: Dict from parse_rqs()
        output_path: Where to save the generated report
        replacements: Optional dict of {old_text: new_text} for custom replacements
    """
    doc = DocxDocument(template_path)

    # Default field mapping: what to find in template -> what to replace with
    # The user can also provide custom replacements
    field_map = {}

    if rqs_data.get("test_regulation"):
        field_map["Test performed according to the Stellantis test procedure B32-3210"] = rqs_data["test_regulation"]

    if rqs_data.get("test_object"):
        field_map["Test applicable on K0 complete front seat"] = rqs_data["test_object"]

    if rqs_data.get("pulse_id"):
        field_map["J050271"] = rqs_data["pulse_id"]

    if rqs_data.get("direction_of_acceleration"):
        field_map["Forward direction"] = rqs_data["direction_of_acceleration"]

    if rqs_data.get("type_of_dummy"):
        field_map["HIII 50% M"] = rqs_data["type_of_dummy"]
        field_map["50% HIII M"] = rqs_data["type_of_dummy"]

    if rqs_data.get("camera_setup"):
        field_map["1 - Mid Left 90° | 2 - Mid Right 90° |"] = rqs_data["camera_setup"]

    # Add custom replacements
    if replacements:
        field_map.update(replacements)

    # Apply all replacements
    applied = {}
    for old_text, new_text in field_map.items():
        if old_text and new_text and old_text != new_text:
            result = _replace_in_document(doc, old_text, new_text)
            applied[old_text] = {"new": new_text, "success": result}

    # Save
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)

    return applied


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 4:
        print("Usage: python report_generator.py <template.docx> <rqs.docx> <output.docx>")
        sys.exit(1)

    rqs_data = parse_rqs(sys.argv[2])
    result = generate_report(sys.argv[1], rqs_data, sys.argv[3])
    print("Replacements applied:")
    for old, info in result.items():
        status = "OK" if info["success"] else "NOT FOUND"
        print(f"  [{status}] '{old[:50]}...' -> '{info['new'][:50]}...'")
