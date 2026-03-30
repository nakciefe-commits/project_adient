"""
Merge taslak1-4.docx files + SPUL graph template into one combined Template.docx
Output: spul/Template.docx
"""
import os
import shutil
from docxcompose.composer import Composer
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy
import re

BASE = os.path.dirname(os.path.abspath(__file__))

taslak_paths = [
    os.path.join(BASE, "taslak1.docx"),
    os.path.join(BASE, "taslak2.docx"),
    os.path.join(BASE, "taslak3.docx"),
    os.path.join(BASE, "taslak4.docx"),
]
spul_template_path = os.path.join(BASE, "spul", "Template.docx")
output_path = os.path.join(BASE, "spul", "Template.docx")
backup_path = os.path.join(BASE, "spul", "Template_backup.docx")

# Backup existing template
if os.path.exists(spul_template_path):
    shutil.copy2(spul_template_path, backup_path)
    print(f"Backed up existing template to: {backup_path}")

# Step 1: Merge taslak1-4 using docxcompose
print("Merging taslak1-4...")
master = Document(taslak_paths[0])
composer = Composer(master)

for path in taslak_paths[1:]:
    doc = Document(path)
    composer.append(doc)

combined_path = os.path.join(BASE, "_combined_taslaks.docx")
composer.save(combined_path)
print(f"Combined taslaks saved to: {combined_path}")

# Step 2: Append SPUL graph pages to combined document
print("Appending SPUL graph pages...")

# Open combined as the base for final output
combined_doc = Document(combined_path)
spul_doc = Document(spul_template_path)

# Add a page break at the end of combined content before graphs
combined_doc.add_page_break()

# Copy body elements from SPUL template (skip the last sectPr)
spul_body = spul_doc.element.body
combined_body = combined_doc.element.body

# Get the sectPr of combined (last element in body)
combined_sectPr = combined_body.find(qn('w:sectPr'))

# Copy relationships (images etc.) from spul_doc to combined_doc
# We need to remap rId references to avoid conflicts
spul_part = spul_doc.part
combined_part = combined_doc.part

# Build rId mapping: spul rId -> new rId in combined
rId_map = {}
for rel_id, rel in spul_part.rels.items():
    if "image" in rel.reltype:
        # Add image to combined document
        image_part = rel.target_part
        new_rId = combined_part.relate_to(image_part, rel.reltype)
        rId_map[rel_id] = new_rId
        print(f"  Image relationship: {rel_id} -> {new_rId}")

# Deep copy SPUL body elements and remap rIds
for child in list(spul_body):
    if child.tag == qn('w:sectPr'):
        continue  # skip sectPr from spul

    elem = copy.deepcopy(child)

    # Remap r:embed attributes in drawings (image references)
    for blip in elem.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        old_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if old_embed and old_embed in rId_map:
            blip.set(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                rId_map[old_embed]
            )

    # Insert before sectPr
    if combined_sectPr is not None:
        combined_body.insert(list(combined_body).index(combined_sectPr), elem)
    else:
        combined_body.append(elem)

# Save final output
combined_doc.save(output_path)
print(f"\nFinal combined template saved to: {output_path}")

# Cleanup temp file
os.remove(combined_path)
print("Done!")
